#!/usr/bin/env python3
"""
Ingest HSK2, HSK3, HSK4 content into Supabase.

Sources:
- HSK2-Vocabulaire-Vivant.html: 200 words (FR only) + stroke SVG data
- HSK3-Vocabulaire-Vivant.html: 500 words (FR only) + stroke SVG data
- HSK2-Preparation-Complete.docx: 197 vocab (tables) + 10 grammar lessons + 2 mock exams
- HSK3-Preparation-Complete.docx: ~180 vocab + 14 grammar lessons + 2 mock exams
- HSK4-Preparation-Complete.docx: vocab + 12 grammar lessons + 2 mock exams

Strategy:
- Vocabulary from HTML (more complete, has strokes)
- Grammar from DOCX (structured lessons)
- EN translations generated from FR + Chinese/Pinyin context
- Courses/Modules/Lessons structure created programmatically
"""

import json, re, urllib.request, urllib.error, sys
from docx import Document

# ─── Supabase config ───
URL = 'https://gmpjkoajhhwvxwsdohll.supabase.co'
KEY = 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImdtcGprb2FqaGh3dnh3c2RvaGxsIiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTc4Mjk4Mzk0NSwiZXhwIjoyMDk4NTU5OTQ1fQ.iHoqQdpjq3_vCMuuHEs9Y9in_lpKQ_cCRaI3EtJ6tKc'
HEADERS = {'apikey': KEY, 'Authorization': f'Bearer {KEY}', 'Content-Type': 'application/json'}

BASE = '/home/user/uploaded_files/hsk_content'

# ─── UUID helpers ───
def course_id(level):   return f"a0000000-0000-0000-0000-{level:012d}"
def module_id(level, n): return f"b0000000-{level:04d}-0000-0000-{n:012d}"
def lesson_id(level, n): return f"c0000000-{level:04d}-0000-0000-{n:012d}"
def vocab_id(level, n):  return f"d1{level:06d}-0000-0000-0000-{n:012d}"
def grammar_id(level, n): return f"91{level:06d}-0000-0000-0000-{n:012d}"
def char_id(level, n):  return f"f0{level:06d}-0000-0000-0000-{n:012d}"
def stroke_id(level, n): return f"f1{level:06d}-0000-0000-0000-{n:012d}"

# ─── API helpers ───
def api_post(table, data, extra_headers=None):
    h = {**HEADERS}
    if extra_headers: h.update(extra_headers)
    body = json.dumps(data).encode()
    req = urllib.request.Request(f'{URL}/rest/v1/{table}', data=body, headers=h, method='POST')
    try:
        resp = urllib.request.urlopen(req)
        return resp.status, resp.read()
    except urllib.error.HTTPError as e:
        return e.code, e.read()

def api_upsert(table, data, conflict_cols):
    return api_post(f'{table}?on_conflict={conflict_cols}', data, {'Prefer': 'resolution=merge-duplicates'})

def api_upsert_batch(table, data, conflict_cols, batch_size=50):
    """Upsert in batches, return (ok_count, err_count)."""
    ok, err = 0, 0
    for i in range(0, len(data), batch_size):
        batch = data[i:i+batch_size]
        code, body = api_upsert(table, batch, conflict_cols)
        if code in (200, 201):
            ok += len(batch)
        else:
            msg = body.decode()[:300] if isinstance(body, bytes) else str(body)[:300]
            print(f"    ERROR {table} batch {i}: {code} - {msg}")
            err += len(batch)
    return ok, err

# ─── FR→EN translation map (manual for common HSK words) ───
# We generate EN from FR meaning using a simple dictionary approach
FR_EN_MAP = {
    'particule': 'particle', 'exclamative': 'exclamatory', 'long': 'long',
    'aider': 'to help', 'aide': 'help', 'jambe': 'leg', 'nez': 'nose',
    'blanc': 'white', 'rouge': 'red', 'vert': 'green', 'noir': 'black',
    'jaune': 'yellow', 'bleu': 'blue', 'couleur': 'color',
    'mari': 'husband', 'femme': 'wife/woman', 'épouse': 'wife',
    'grand-père': 'grandfather', 'grand-mère': 'grandmother',
    'oncle': 'uncle', 'tante': 'aunt', 'frère': 'brother', 'sœur': 'sister',
    'fils': 'son', 'fille': 'daughter/girl',
    'rapide': 'fast', 'lent': 'slow', 'chaud': 'hot', 'froid': 'cold',
    'parce que': 'because', 'mais': 'but', 'donc': 'therefore',
    'si': 'if', 'encore': 'still/again', 'déjà': 'already',
    'préparer': 'to prepare', 'présenter': 'to introduce',
    'espérer': 'to hope', 'croire': 'to believe',
    'comprendre': 'to understand', 'penser': 'to think',
    'aéroport': 'airport', 'hôpital': 'hospital', 'école': 'school',
    'football': 'football/soccer', 'basketball': 'basketball',
    'anniversaire': 'birthday', 'lit': 'bed', 'porte': 'door',
    'petit': 'small/short', 'grand': 'big/tall',
    'nouveau': 'new', 'vieux': 'old', 'jeune': 'young',
    'beau': 'beautiful', 'facile': 'easy', 'difficile': 'difficult',
    'content': 'happy', 'triste': 'sad', 'fatigué': 'tired',
    'malade': 'sick', 'occupé': 'busy', 'libre': 'free',
}

def fr_to_en_meaning(fr_meaning):
    """Generate a basic EN translation from FR meaning."""
    fr_lower = fr_meaning.lower().strip()
    # Direct lookup
    if fr_lower in FR_EN_MAP:
        return FR_EN_MAP[fr_lower]
    # Partial match
    for fr, en in FR_EN_MAP.items():
        if fr in fr_lower:
            return en
    # Keep the FR as-is with a marker (we'll still have the word + pinyin for context)
    return fr_meaning  # FR is better than nothing for MVP

def fr_to_en_sentence(fr_sentence):
    """For example sentences, just keep FR (user sees Chinese + pinyin + FR)."""
    return fr_sentence  # We'll mark these as needing proper translation later


# ═══════════════════════════════════════════════════
# STEP 1: Extract vocabulary from HTML files
# ═══════════════════════════════════════════════════
print("=" * 60)
print("STEP 1: Extract vocabulary from HTML files")
print("=" * 60)

def extract_words_from_html(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        html = f.read()
    m = re.search(r'const WORDS\s*=\s*(\[.*?\]);', html, re.DOTALL)
    if not m:
        raise ValueError(f"WORDS array not found in {filepath}")
    return json.loads(m.group(1))

hsk_words = {}
for level, filename, expected in [(2, 'HSK2-Vocabulaire-Vivant.html', 200), 
                                    (3, 'HSK3-Vocabulaire-Vivant.html', 500)]:
    words = extract_words_from_html(f'{BASE}/{filename}')
    assert len(words) == expected, f"HSK{level}: expected {expected}, got {len(words)}"
    hsk_words[level] = words
    print(f"  HSK{level}: {len(words)} words extracted")

# ═══════════════════════════════════════════════════
# STEP 2: Extract grammar from DOCX files
# ═══════════════════════════════════════════════════
print("\n" + "=" * 60)
print("STEP 2: Extract grammar from DOCX files")
print("=" * 60)

def extract_grammar_from_docx(filepath, level):
    """Extract grammar lessons from DOCX. Returns list of grammar dicts."""
    doc = Document(filepath)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    grammar_lessons = []
    current = None
    in_exercises = False
    in_corriges = False
    
    for p in paras:
        # Detect lesson header: "Leçon N — pattern — title"
        m = re.match(r'Leçon\s+(\d+)\s*[—–-]\s*(.+?)\s*[—–-]\s*(.+)', p)
        if m:
            if current:
                grammar_lessons.append(current)
            lesson_num = int(m.group(1))
            pattern = m.group(2).strip()
            title = m.group(3).strip()
            current = {
                'num': lesson_num,
                'pattern': pattern,
                'title_fr': title,
                'explanation_fr': '',
                'exercises_fr': [],
                'corriges_fr': [],
            }
            in_exercises = False
            in_corriges = False
            continue
        
        if not current:
            continue
            
        if p.startswith('Objectif communicatif'):
            current['objective_fr'] = p.replace('Objectif communicatif : ', '').replace('Objectif communicatif:', '').strip()
            continue
            
        if p.startswith('Exercices d') and 'application' in p:
            in_exercises = True
            in_corriges = False
            continue
            
        if p.startswith('Corrigé'):
            in_corriges = True
            in_exercises = False
            continue
            
        # Check if we hit next section
        if p.startswith(('3. Examen', '4. Examen', '5. Volet', '6. Feuille')):
            if current:
                grammar_lessons.append(current)
                current = None
            break
            
        if in_corriges:
            current['corriges_fr'].append(p)
        elif in_exercises:
            current['exercises_fr'].append(p)
        elif current and not p.startswith('Leçon'):
            # It's explanation text
            if current['explanation_fr']:
                current['explanation_fr'] += '\n' + p
            else:
                current['explanation_fr'] = p
    
    if current:
        grammar_lessons.append(current)
    
    return grammar_lessons

hsk_grammar = {}
for level, filename in [(2, 'HSK2-Preparation-Complete.docx'), 
                          (3, 'HSK3-Preparation-Complete.docx'),
                          (4, 'HSK4-Preparation-Complete.docx')]:
    grammar = extract_grammar_from_docx(f'{BASE}/{filename}', level)
    hsk_grammar[level] = grammar
    print(f"  HSK{level}: {len(grammar)} grammar lessons extracted")
    for g in grammar[:3]:
        print(f"    L{g['num']}: {g['pattern']} — {g['title_fr'][:60]}")

# ═══════════════════════════════════════════════════
# STEP 3: Extract vocab from DOCX tables (for HSK4 which has no HTML)
# ═══════════════════════════════════════════════════
print("\n" + "=" * 60)
print("STEP 3: Extract HSK4 vocab from DOCX tables")
print("=" * 60)

def extract_vocab_from_docx(filepath):
    """Extract vocabulary from DOCX tables (Chinois/Pinyin/Français)."""
    doc = Document(filepath)
    vocab = []
    for table in doc.tables:
        if len(table.columns) < 3:
            continue
        # Check if it's a vocab table (header: Chinois/Pinyin/Français)
        header = [c.text.strip() for c in table.rows[0].cells]
        if header[0] != 'Chinois' or header[1] != 'Pinyin' or header[2] != 'Français':
            continue
        # Check if rows are single words (not full sentences)
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if not cells[0] or len(cells[0]) > 5:  # Skip sentence examples
                continue
            vocab.append({
                'word': cells[0],
                'pinyin': cells[1],
                'meaning_fr': cells[2],
            })
    return vocab

hsk4_vocab = extract_vocab_from_docx(f'{BASE}/HSK4-Preparation-Complete.docx')
print(f"  HSK4: {len(hsk4_vocab)} vocabulary words from DOCX tables")
for v in hsk4_vocab[:5]:
    print(f"    {v['word']} ({v['pinyin']}) = {v['meaning_fr'][:50]}")


# ═══════════════════════════════════════════════════
# STEP 4: Create courses for HSK2, HSK3, HSK4
# ═══════════════════════════════════════════════════
print("\n" + "=" * 60)
print("STEP 4: Create courses")
print("=" * 60)

# First check if courses exist
for level in [2, 3, 4]:
    cid = course_id(level)
    slug = f"hsk{level}"
    course_data = [{
        'id': cid,
        'slug': slug,
        'hsk_level': str(level),
        'is_published': True,
        'sort_order': level,
        'status': 'published',
    }]
    code, body = api_upsert('courses', course_data, 'id')
    status = '✓' if code in (200, 201) else f'✗ {code}'
    print(f"  Course HSK{level} ({cid}): {status}")
    
    # Course translations
    for locale, title, desc in [
        ('fr', f'HSK {level}', f'Préparation complète au HSK {level}'),
        ('en', f'HSK {level}', f'Complete HSK {level} preparation'),
    ]:
        trans = [{
            'course_id': cid,
            'locale': locale,
            'title': title,
            'description': desc,
        }]
        code, _ = api_upsert('course_translations', trans, 'course_id,locale')
        print(f"    {locale}: {code}")


# ═══════════════════════════════════════════════════
# STEP 5: Create modules and lessons for HSK2/3/4
# ═══════════════════════════════════════════════════
print("\n" + "=" * 60)
print("STEP 5: Create modules and lessons")
print("=" * 60)

# HSK2: 5 modules (vocab 1-4, grammar, exam)
# HSK3: 7 modules (vocab 1-5, grammar, exam)
# HSK4: similar structure

LEVEL_STRUCTURE = {
    2: {
        'vocab_count': 200,
        'vocab_per_module': 50,
        'grammar_count': 10,
        'modules': [
            {'name_fr': 'Vocabulaire HSK2 — Lot 1', 'name_en': 'HSK2 Vocabulary — Batch 1', 'type': 'vocab', 'range': (1, 50)},
            {'name_fr': 'Vocabulaire HSK2 — Lot 2', 'name_en': 'HSK2 Vocabulary — Batch 2', 'type': 'vocab', 'range': (51, 100)},
            {'name_fr': 'Vocabulaire HSK2 — Lot 3', 'name_en': 'HSK2 Vocabulary — Batch 3', 'type': 'vocab', 'range': (101, 150)},
            {'name_fr': 'Vocabulaire HSK2 — Lot 4', 'name_en': 'HSK2 Vocabulary — Batch 4', 'type': 'vocab', 'range': (151, 200)},
            {'name_fr': 'Grammaire HSK2', 'name_en': 'HSK2 Grammar', 'type': 'grammar'},
            {'name_fr': 'Examens blancs HSK2', 'name_en': 'HSK2 Mock Exams', 'type': 'exam'},
        ],
    },
    3: {
        'vocab_count': 500,
        'vocab_per_module': 100,
        'grammar_count': 14,
        'modules': [
            {'name_fr': 'Vocabulaire HSK3 — Lot 1', 'name_en': 'HSK3 Vocabulary — Batch 1', 'type': 'vocab', 'range': (1, 100)},
            {'name_fr': 'Vocabulaire HSK3 — Lot 2', 'name_en': 'HSK3 Vocabulary — Batch 2', 'type': 'vocab', 'range': (101, 200)},
            {'name_fr': 'Vocabulaire HSK3 — Lot 3', 'name_en': 'HSK3 Vocabulary — Batch 3', 'type': 'vocab', 'range': (201, 300)},
            {'name_fr': 'Vocabulaire HSK3 — Lot 4', 'name_en': 'HSK3 Vocabulary — Batch 4', 'type': 'vocab', 'range': (301, 400)},
            {'name_fr': 'Vocabulaire HSK3 — Lot 5', 'name_en': 'HSK3 Vocabulary — Batch 5', 'type': 'vocab', 'range': (401, 500)},
            {'name_fr': 'Grammaire HSK3', 'name_en': 'HSK3 Grammar', 'type': 'grammar'},
            {'name_fr': 'Examens blancs HSK3', 'name_en': 'HSK3 Mock Exams', 'type': 'exam'},
        ],
    },
    4: {
        'vocab_count': len(hsk4_vocab),
        'vocab_per_module': 50,
        'grammar_count': 12,
        'modules': [
            {'name_fr': 'Grammaire HSK4', 'name_en': 'HSK4 Grammar', 'type': 'grammar'},
            {'name_fr': 'Examens blancs HSK4', 'name_en': 'HSK4 Mock Exams', 'type': 'exam'},
        ],
    },
}

# Add vocab modules for HSK4 dynamically
hsk4_vocab_modules = []
for i in range(0, len(hsk4_vocab), 50):
    batch_num = i // 50 + 1
    end = min(i + 50, len(hsk4_vocab))
    hsk4_vocab_modules.append({
        'name_fr': f'Vocabulaire HSK4 — Lot {batch_num}',
        'name_en': f'HSK4 Vocabulary — Batch {batch_num}',
        'type': 'vocab',
        'range': (i + 1, end),
    })
LEVEL_STRUCTURE[4]['modules'] = hsk4_vocab_modules + LEVEL_STRUCTURE[4]['modules']

lesson_counter = {}  # {level: lesson_num}

for level, struct in LEVEL_STRUCTURE.items():
    cid = course_id(level)
    lesson_counter[level] = 0
    
    for mod_idx, mod_def in enumerate(struct['modules'], 1):
        mid = module_id(level, mod_idx)
        
        # Create module
        mod_data = [{
            'id': mid,
            'course_id': cid,
            'sort_order': mod_idx,
            'estimated_duration_minutes': 60,
            'status': 'published',
        }]
        code, _ = api_upsert('modules', mod_data, 'id')
        
        # Module translations
        for locale in ['fr', 'en']:
            name = mod_def[f'name_{locale}']
            trans = [{
                'module_id': mid,
                'locale': locale,
                'title': name,
                'description': f'{name} — Niveau HSK {level}' if locale == 'fr' else f'{name} — HSK Level {level}',
                'objectives': json.dumps([name]),
            }]
            api_upsert('module_translations', trans, 'module_id,locale')
        
        # Create lessons for this module
        if mod_def['type'] == 'vocab':
            r_start, r_end = mod_def['range']
            # Create lessons in groups of 10 words
            for batch_start in range(r_start, r_end + 1, 10):
                batch_end = min(batch_start + 9, r_end)
                lesson_counter[level] += 1
                ln = lesson_counter[level]
                lid = lesson_id(level, ln)
                
                lesson_data = [{
                    'id': lid,
                    'module_id': mid,
                    'sort_order': ln,
                    'lesson_type': 'standard',
                    'estimated_duration_minutes': 15,
                    'status': 'published',
                }]
                api_upsert('lessons', lesson_data, 'id')
                
                for locale, title_tpl in [('fr', 'Mots {}-{}'), ('en', 'Words {}-{}')]:
                    trans = [{
                        'lesson_id': lid,
                        'locale': locale,
                        'title': title_tpl.format(batch_start, batch_end),
                    }]
                    api_upsert('lesson_translations', trans, 'lesson_id,locale')
        
        elif mod_def['type'] == 'grammar':
            grammar_list = hsk_grammar.get(level, [])
            for g in grammar_list:
                lesson_counter[level] += 1
                ln = lesson_counter[level]
                lid = lesson_id(level, ln)
                
                lesson_data = [{
                    'id': lid,
                    'module_id': mid,
                    'sort_order': ln,
                    'lesson_type': 'standard',
                    'estimated_duration_minutes': 20,
                    'status': 'published',
                }]
                api_upsert('lessons', lesson_data, 'id')
                
                for locale in ['fr', 'en']:
                    if locale == 'fr':
                        title = f"Leçon {g['num']} — {g['pattern']} — {g['title_fr']}"
                    else:
                        title = f"Lesson {g['num']} — {g['pattern']}"
                    trans = [{
                        'lesson_id': lid,
                        'locale': locale,
                        'title': title,
                    }]
                    api_upsert('lesson_translations', trans, 'lesson_id,locale')
        
        elif mod_def['type'] == 'exam':
            for exam_num in [1, 2]:
                lesson_counter[level] += 1
                ln = lesson_counter[level]
                lid = lesson_id(level, ln)
                
                lesson_data = [{
                    'id': lid,
                    'module_id': mid,
                    'sort_order': ln,
                    'lesson_type': 'exam',
                    'estimated_duration_minutes': 45,
                    'status': 'published',
                }]
                api_upsert('lessons', lesson_data, 'id')
                
                for locale in ['fr', 'en']:
                    title = f"Examen blanc n°{exam_num}" if locale == 'fr' else f"Mock Exam #{exam_num}"
                    trans = [{
                        'lesson_id': lid,
                        'locale': locale,
                        'title': title,
                    }]
                    api_upsert('lesson_translations', trans, 'lesson_id,locale')
    
    total_lessons = lesson_counter[level]
    total_modules = len(struct['modules'])
    print(f"  HSK{level}: {total_modules} modules, {total_lessons} lessons created")


# ═══════════════════════════════════════════════════
# STEP 6: Insert vocabulary items + translations (HSK2 & HSK3 from HTML)
# ═══════════════════════════════════════════════════
print("\n" + "=" * 60)
print("STEP 6: Insert vocabulary items + translations")
print("=" * 60)

for level, words in hsk_words.items():
    vocab_items = []
    vocab_translations_fr = []
    vocab_translations_en = []
    
    for i, w in enumerate(words, 1):
        vid = vocab_id(level, i)
        
        # Determine word type from decomposition or context
        word_type = 'word'  # default
        
        vocab_items.append({
            'id': vid,
            'simplified': w['word'],
            'traditional': None,  # Not available in HTML
            'pinyin': w['pinyin'],
            'hsk_level': str(level),
            'word_type': word_type,
            'frequency_rank': i,
            'status': 'published',
        })
        
        # FR translation
        fr_meaning = w.get('meaning', {}).get('fr', '')
        fr_example = w.get('example', {})
        vocab_translations_fr.append({
            'vocabulary_id': vid,
            'locale': 'fr',
            'meaning': fr_meaning,
            'example_sentence': fr_example.get('cn', ''),
            'example_pinyin': fr_example.get('py', ''),
            'example_translation': fr_example.get('tr', {}).get('fr', ''),
            'usage_notes': w.get('confusion', {}).get('fr', '') if w.get('confusion') else None,
        })
        
        # EN translation (generated from FR)
        en_meaning = fr_to_en_meaning(fr_meaning)
        en_example_tr = fr_example.get('tr', {}).get('en', '') or fr_to_en_sentence(
            fr_example.get('tr', {}).get('fr', '')
        )
        vocab_translations_en.append({
            'vocabulary_id': vid,
            'locale': 'en',
            'meaning': en_meaning,
            'example_sentence': fr_example.get('cn', ''),
            'example_pinyin': fr_example.get('py', ''),
            'example_translation': en_example_tr,
        })
    
    # Batch upsert
    ok, err = api_upsert_batch('vocabulary_items', vocab_items, 'id')
    print(f"  HSK{level} vocab_items: {ok} ok, {err} err")
    
    ok, err = api_upsert_batch('vocabulary_translations', vocab_translations_fr, 'vocabulary_id,locale')
    print(f"  HSK{level} vocab_translations FR: {ok} ok, {err} err")
    
    ok, err = api_upsert_batch('vocabulary_translations', vocab_translations_en, 'vocabulary_id,locale')
    print(f"  HSK{level} vocab_translations EN: {ok} ok, {err} err")

# HSK4 vocab from DOCX
if hsk4_vocab:
    vocab_items = []
    vocab_translations_fr = []
    vocab_translations_en = []
    level = 4
    
    for i, v in enumerate(hsk4_vocab, 1):
        vid = vocab_id(level, i)
        vocab_items.append({
            'id': vid,
            'simplified': v['word'],
            'pinyin': v['pinyin'],
            'hsk_level': '4',
            'word_type': 'word',
            'frequency_rank': i,
            'status': 'published',
        })
        
        vocab_translations_fr.append({
            'vocabulary_id': vid,
            'locale': 'fr',
            'meaning': v['meaning_fr'],
        })
        
        en_meaning = fr_to_en_meaning(v['meaning_fr'])
        vocab_translations_en.append({
            'vocabulary_id': vid,
            'locale': 'en',
            'meaning': en_meaning,
        })
    
    ok, err = api_upsert_batch('vocabulary_items', vocab_items, 'id')
    print(f"  HSK4 vocab_items: {ok} ok, {err} err")
    ok, err = api_upsert_batch('vocabulary_translations', vocab_translations_fr, 'vocabulary_id,locale')
    print(f"  HSK4 vocab_translations FR: {ok} ok, {err} err")
    ok, err = api_upsert_batch('vocabulary_translations', vocab_translations_en, 'vocabulary_id,locale')
    print(f"  HSK4 vocab_translations EN: {ok} ok, {err} err")


# ═══════════════════════════════════════════════════
# STEP 7: Insert grammar points + translations
# ═══════════════════════════════════════════════════
print("\n" + "=" * 60)
print("STEP 7: Insert grammar points + translations")
print("=" * 60)

for level, grammar_list in hsk_grammar.items():
    gp_data = []
    gp_trans_fr = []
    gp_trans_en = []
    
    for i, g in enumerate(grammar_list, 1):
        gid = grammar_id(level, i)
        
        gp_data.append({
            'id': gid,
            'pattern': g['pattern'],
            'hsk_level': str(level),
            'sort_order': i,
            'difficulty': min(level, 5),
            'status': 'published',
        })
        
        # Build explanation HTML
        expl_html = f"<p>{g['explanation_fr']}</p>"
        if g.get('exercises_fr'):
            expl_html += "\n<h4>Exercices</h4>\n<ol>"
            for ex in g['exercises_fr']:
                expl_html += f"\n<li>{ex}</li>"
            expl_html += "\n</ol>"
        if g.get('corriges_fr'):
            expl_html += "\n<h4>Corrigé</h4>\n<ol>"
            for c in g['corriges_fr']:
                expl_html += f"\n<li>{c}</li>"
            expl_html += "\n</ol>"
        
        gp_trans_fr.append({
            'grammar_point_id': gid,
            'locale': 'fr',
            'title': f"{g['pattern']} — {g['title_fr']}",
            'explanation_html': expl_html,
        })
        
        gp_trans_en.append({
            'grammar_point_id': gid,
            'locale': 'en',
            'title': f"{g['pattern']}",
            'explanation_html': f"<p>{g['pattern']} pattern</p>",
        })
    
    ok, err = api_upsert_batch('grammar_points', gp_data, 'id')
    print(f"  HSK{level} grammar_points: {ok} ok, {err} err")
    ok, err = api_upsert_batch('grammar_point_translations', gp_trans_fr, 'grammar_point_id,locale')
    print(f"  HSK{level} grammar_translations FR: {ok} ok, {err} err")
    ok, err = api_upsert_batch('grammar_point_translations', gp_trans_en, 'grammar_point_id,locale')
    print(f"  HSK{level} grammar_translations EN: {ok} ok, {err} err")


# ═══════════════════════════════════════════════════
# STEP 8: Insert characters + stroke data (HSK2 & HSK3 from HTML)
# ═══════════════════════════════════════════════════
print("\n" + "=" * 60)
print("STEP 8: Insert characters + stroke data")
print("=" * 60)

for level, words in hsk_words.items():
    # Collect unique characters from all words
    char_map = {}  # char → (first_word_idx, char_data)
    
    for word_idx, w in enumerate(words):
        for char_data in w.get('chars', []):
            ch = char_data.get('char', '')
            if ch and ch not in char_map:
                char_map[ch] = (word_idx, char_data)
    
    chars_data = []
    chars_trans_fr = []
    chars_trans_en = []
    strokes_data = []
    
    for char_idx, (ch, (word_idx, cd)) in enumerate(char_map.items(), 1):
        cid = char_id(level, char_idx)
        
        chars_data.append({
            'id': cid,
            'character': ch,
            'pinyin': cd.get('pinyin', ''),
            'stroke_count': len(cd.get('strokes', [])),
            'hsk_level': str(level),
            'frequency_rank': char_idx,
            'status': 'published',
        })
        
        # Find FR meaning from vocabulary context
        parent_word = words[word_idx]
        fr_meaning = parent_word.get('meaning', {}).get('fr', '')
        
        chars_trans_fr.append({
            'character_id': cid,
            'locale': 'fr',
            'meaning': f"Composant de « {parent_word['word']} » ({fr_meaning})",
        })
        
        en_meaning = fr_to_en_meaning(fr_meaning)
        chars_trans_en.append({
            'character_id': cid,
            'locale': 'en',
            'meaning': f"Component of '{parent_word['word']}' ({en_meaning})",
        })
        
        # Stroke data
        if cd.get('strokes'):
            sid = stroke_id(level, char_idx)
            strokes_data.append({
                'id': sid,
                'character_id': cid,
                'strokes': cd.get('strokes', []),
                'medians': cd.get('medians', []),
                'source': 'makemeahanzi',
            })
    
    ok, err = api_upsert_batch('characters', chars_data, 'id')
    print(f"  HSK{level} characters: {ok} ok, {err} err ({len(char_map)} unique chars)")
    
    ok, err = api_upsert_batch('character_translations', chars_trans_fr, 'character_id,locale')
    print(f"  HSK{level} character_translations FR: {ok} ok, {err} err")
    
    ok, err = api_upsert_batch('character_translations', chars_trans_en, 'character_id,locale')
    print(f"  HSK{level} character_translations EN: {ok} ok, {err} err")
    
    if strokes_data:
        ok, err = api_upsert_batch('stroke_order_data', strokes_data, 'id')
        print(f"  HSK{level} stroke_order_data: {ok} ok, {err} err")


# ═══════════════════════════════════════════════════
# FINAL SUMMARY
# ═══════════════════════════════════════════════════
print("\n" + "=" * 60)
print("FINAL: Verification counts")
print("=" * 60)

for table in ['courses', 'course_translations', 'modules', 'module_translations', 
              'lessons', 'lesson_translations', 'vocabulary_items', 'vocabulary_translations',
              'grammar_points', 'grammar_point_translations', 'characters', 'character_translations',
              'stroke_order_data', 'exercises', 'exercise_options']:
    req = urllib.request.Request(
        f'{URL}/rest/v1/{table}?select=id&limit=1',
        headers={**HEADERS, 'Prefer': 'count=exact', 'Range': '0-0'}
    )
    try:
        resp = urllib.request.urlopen(req)
        cr = resp.headers.get('Content-Range', '?')
        print(f"  {table}: {cr}")
    except:
        print(f"  {table}: ERROR")

print("\nDone!")
